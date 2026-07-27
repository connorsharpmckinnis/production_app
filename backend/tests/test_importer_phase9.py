"""Phase 9 importer grammar, DOCX, parity, and diagnostics tests."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from app.db.seed import seed_database
from app.models import Act, Character, Moment, Organization, Production, Song
from app.services.importer import ImportLineError, import_script
from app.services.importer.formats.docx import extract_docx_lines
from app.services.importer.grammar import (
    SpeakerListError,
    is_all_caps_lyric,
    parse_speaker_list,
    validate_song_title,
)
from tests.helpers.import_fingerprint import semantic_import_fingerprint


FIXTURES = Path(__file__).resolve().parents[2] / "fixtures" / "scripts"


@pytest.fixture
def seeded_db(db_session: Session, test_settings) -> Session:
    seed_database(db_session, test_settings)
    return db_session


@pytest.fixture
def production(seeded_db: Session) -> Production:
    organization = seeded_db.query(Organization).first()
    production = Production(organization_id=organization.id, title="Phase 9")
    seeded_db.add(production)
    seeded_db.commit()
    return production


@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        ("MS. ELEPHANT", ["MS. ELEPHANT"]),
        ("BRI'ISH NEWSIE", ["BRI'ISH NEWSIE"]),
        ("ORDE-LEES", ["ORDE-LEES"]),
        ("CREW 2", ["CREW 2"]),
        ("MS. ELEPHANT, ORDE-LEES", ["MS. ELEPHANT", "ORDE-LEES"]),
        ("VERA & MOM", ["VERA", "MOM"]),
        (
            "SHACKLETON, WORSLEY, and CREAN",
            ["SHACKLETON", "WORSLEY", "CREAN"],
        ),
    ],
)
def test_parse_speaker_list_accepts_documented_names(raw, expected):
    assert parse_speaker_list(raw) == expected


@pytest.mark.parametrize(
    "raw",
    ["McNISH", ".ELEPHANT", "NAME!", "NAME,, OTHER", "NAME and ", "123"],
)
def test_parse_speaker_list_rejects_invalid_names(raw):
    with pytest.raises(SpeakerListError):
        parse_speaker_list(raw)


@pytest.mark.parametrize(
    "lyric",
    [
        "IS THIS REAL?!",
        "WE GO (TOGETHER)",
        "WAIT—DON'T GO…",
        "“YES,” WE SING.",
        "ONE – TWO & THREE",
        "I'M/HE'S JUST A MAN",
        "LEFT US HERE; ALONE WITH THAT FLOCK",
    ],
)
def test_lyric_validator_accepts_approved_punctuation(lyric):
    assert is_all_caps_lyric(lyric)


@pytest.mark.parametrize(
    "text",
    [
        "Lowercase prose.",
        "SHACKLETON: THIS IS DIALOGUE",
        "UNBALANCED (LYRIC",
        "NOT@APPROVED",
        "AMEN[^9]",
    ],
)
def test_lyric_validator_rejects_non_lyrics(text):
    assert not is_all_caps_lyric(text)


def test_validate_song_title_rejects_performer_shaped_headers():
    with pytest.raises(ValueError, match="performer attribution"):
        validate_song_title("ALL", {"ALL", "ENSEMBLE", "SHACKLETON"})


def test_punctuated_group_dialogue_creates_deduplicated_characters(
    seeded_db,
    production,
):
    script = """# Act One
## Scene One - Test
MS. ELEPHANT, ORDE-LEES: Together.
MS. ELEPHANT: Again.
"""
    result = import_script(seeded_db, production, script)

    assert result.moments_created == 2
    names = {
        character.name
        for character in seeded_db.query(Character).order_by(Character.name)
    }
    assert names == {"ALL", "ENSEMBLE", "MS. ELEPHANT", "ORDE-LEES"}


def test_punctuated_dialogue_speaker_can_be_song_performer(seeded_db, production):
    script = """# Act One
## Scene One - Test
MS. ELEPHANT: Hello.
### MY SONG
MS. ELEPHANT
LINE ONE HERE
"""
    import_script(seeded_db, production, script)
    moments = seeded_db.query(Moment).order_by(Moment.sequence_number).all()
    assert [moment.moment_type.name for moment in moments] == [
        "dialogue",
        "song_header",
        "song_attribution",
        "lyric",
    ]


def test_mixed_case_dialogue_fails_with_actionable_error_and_rolls_back(
    seeded_db,
    production,
):
    script = """# Act One
## Scene One - Test
McNISH: Mixed case.
"""
    with pytest.raises(ImportLineError, match="ALL CAPS") as exc:
        import_script(seeded_db, production, script)

    assert exc.value.line_number == 3
    assert exc.value.context_snippet == (
        "# Act One\n## Scene One - Test\nMcNISH: Mixed case."
    )
    assert seeded_db.query(Moment).count() == 0
    assert seeded_db.query(Act).count() == 0
    assert seeded_db.query(Song).count() == 0
    assert seeded_db.query(Character).count() == 0


def test_spaced_dialogue_colon_fails_clearly(seeded_db, production):
    script = """# Act One
## Scene One - Test
APPLICANT : Hudson, Hudson! Are you applying?
"""
    with pytest.raises(ImportLineError, match="no spaces before the colon"):
        import_script(seeded_db, production, script)


def test_song_header_cannot_be_performer_attribution(seeded_db, production):
    script = """# Act One
## Scene One - Test
### ALL
"""
    with pytest.raises(ImportLineError, match="performer attribution") as exc:
        import_script(seeded_db, production, script)

    assert seeded_db.query(Song).count() == 0
    assert seeded_db.query(Moment).count() == 0
    assert exc.value.source_format == "md"


def test_h4_author_note_is_not_swallowed_as_song_description(seeded_db, production):
    script = """# Act One
## Scene One - Test
### MY SONG
#### Note: Keep this as an author note.
"""
    import_script(seeded_db, production, script)
    song = seeded_db.query(Song).one()
    assert song.description is None
    note = seeded_db.query(Moment).order_by(Moment.sequence_number.desc()).first()
    assert note.moment_type.name == "author_note"


def test_moment_before_scene_reports_real_line_number(seeded_db, production):
    script = """# Act One
*Lights up before any scene.*
"""
    with pytest.raises(ImportLineError) as exc:
        import_script(seeded_db, production, script)
    assert exc.value.line_number == 2
    assert "No Scene defined yet" in exc.value.message


def test_punctuated_lyrics_preserve_text(seeded_db, production):
    lyric = "WAIT—DON'T GO (YET)… IS THIS REAL?!"
    script = f"""# Act One
## Scene One - Test
### MY SONG
ALL
{lyric}
"""
    import_script(seeded_db, production, script)

    moment = seeded_db.query(Moment).order_by(Moment.sequence_number.desc()).first()
    assert moment.original_text == lyric


def test_mixed_case_h4_marker_fails_clearly(seeded_db, production):
    script = """# Act One
## Scene One - Test
### MY SONG
#### rCHORUS
"""
    with pytest.raises(ImportLineError, match="mixed-case song marker"):
        import_script(seeded_db, production, script)


def test_dry_run_does_not_persist_import(seeded_db, production):
    script = """# Act One
## Scene One - Test
*Lights up.*
"""
    result = import_script(seeded_db, production, script, dry_run=True)
    assert result.moments_created == 1
    assert seeded_db.query(Act).count() == 0
    assert seeded_db.query(Moment).count() == 0


def _append_modern_hyperlink_text(paragraph, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_modern_hyperlinked_heading_3_preserves_visible_song_title(
    seeded_db,
    production,
):
    document = Document()
    document.add_heading("Act One", level=1)
    document.add_heading("Scene One - Test", level=2)
    song_heading = document.add_heading("", level=3)
    _append_modern_hyperlink_text(song_heading, "INTO THE DEEP (PRE-PRISE)")
    song_heading.add_run("")
    document.add_paragraph("ALL")
    document.add_paragraph("WE GOOOO… INTO THE DEEP")
    content = BytesIO()
    document.save(content)

    lines = extract_docx_lines(content.getvalue())
    assert "### INTO THE DEEP (PRE-PRISE)" in lines

    result = import_script(
        seeded_db,
        production,
        content.getvalue(),
        filename="modern-hyperlink.docx",
    )
    assert result.songs_created == 1
    moments = seeded_db.query(Moment).order_by(Moment.sequence_number).all()
    assert [moment.moment_type.name for moment in moments] == [
        "song_header",
        "song_attribution",
        "lyric",
    ]


def test_docx_heading_3_all_fails_with_paragraph_metadata(seeded_db, production):
    document = Document()
    document.add_heading("Act One", level=1)
    document.add_heading("Scene One - Test", level=2)
    document.add_heading("ALL", level=3)
    content = BytesIO()
    document.save(content)

    with pytest.raises(ImportLineError, match="performer attribution") as exc:
        import_script(
            seeded_db,
            production,
            content.getvalue(),
            filename="heading3-all.docx",
        )

    assert exc.value.source_format == "docx"
    assert exc.value.paragraph_number == 3
    assert exc.value.paragraph_style == "Heading 3"
    assert exc.value.issues[0].kind == "song"
    assert exc.value.context_snippet is not None
    assert exc.value.context_snippet.splitlines()[-1] == "### ALL"
    assert "Act One" in exc.value.context_snippet
    assert "Scene One - Test" in exc.value.context_snippet


def test_scene_one_markdown_and_docx_have_exact_semantic_parity(
    seeded_db,
    production,
):
    organization = seeded_db.query(Organization).first()
    docx_production = Production(
        organization_id=organization.id,
        title="Phase 9 DOCX",
    )
    seeded_db.add(docx_production)
    seeded_db.commit()

    import_script(
        seeded_db,
        production,
        (FIXTURES / "endurance-scene1.md").read_bytes(),
        filename="endurance-scene1.md",
    )
    import_script(
        seeded_db,
        docx_production,
        (FIXTURES / "Endurance Scene 1.docx").read_bytes(),
        filename="Endurance Scene 1.docx",
    )

    assert semantic_import_fingerprint(
        seeded_db,
        production,
    ) == semantic_import_fingerprint(seeded_db, docx_production)


def test_cleaned_full_scripts_import_with_structural_parity(seeded_db, production):
    organization = seeded_db.query(Organization).first()
    docx_production = Production(
        organization_id=organization.id,
        title="Phase 9 Full DOCX",
    )
    seeded_db.add(docx_production)
    seeded_db.commit()

    md_result = import_script(
        seeded_db,
        production,
        (FIXTURES / "endurance-full-cleaned.md").read_bytes(),
        filename="endurance-full-cleaned.md",
    )
    docx_result = import_script(
        seeded_db,
        docx_production,
        (FIXTURES / "endurance-full-cleaned.docx").read_bytes(),
        filename="endurance-full-cleaned.docx",
    )

    assert md_result.acts_created == 2
    assert docx_result.acts_created == 2
    assert md_result.scenes_created == 15
    assert docx_result.scenes_created == 15
    assert md_result.songs_created == 25
    assert docx_result.songs_created == 25

    md_fp = semantic_import_fingerprint(seeded_db, production)
    docx_fp = semantic_import_fingerprint(seeded_db, docx_production)
    assert md_fp["structure"] == docx_fp["structure"]
    assert md_fp["songs"] == docx_fp["songs"]
    assert md_fp["characters"] == docx_fp["characters"]
    # Moment packing differs by export format (scripture line breaks, Heading 4
    # vs Body lyrics). Structural parity above is the Phase 9 full-script bar.
    assert len(md_fp["moments"]) >= 1300
    assert len(docx_fp["moments"]) >= 1300


def test_dirty_full_markdown_still_fails_on_known_source_issue(seeded_db, production):
    with pytest.raises(ImportLineError) as exc:
        import_script(
            seeded_db,
            production,
            (FIXTURES / "endurance-full.md").read_bytes(),
            filename="endurance-full.md",
        )
    assert len(exc.value.issues) >= 1
    assert seeded_db.query(Moment).count() == 0


def test_collects_multiple_issues_and_collapses_song_errors(seeded_db, production):
    script = """# Act One
## Scene One - Test
this prose is not a moment
APPLICANT: Valid dialogue.
McNISH: Bad casing.
### ALL
SKIPPED LYRIC LINE
## Scene Two - Later
also not a moment
"""
    with pytest.raises(ImportLineError) as exc:
        import_script(seeded_db, production, script)

    issues = exc.value.issues
    assert len(issues) == 4
    assert issues[0].kind == "line"
    assert "Unrecognized prose" in issues[0].message
    assert issues[1].kind == "line"
    assert "ALL CAPS" in issues[1].message
    assert issues[2].kind == "song"
    assert "performer attribution" in issues[2].message
    assert issues[3].kind == "line"
    assert "Unrecognized prose" in issues[3].message
    assert issues[3].line_number > issues[2].line_number
    assert seeded_db.query(Moment).count() == 0
    assert seeded_db.query(Act).count() == 0
    assert seeded_db.query(Song).count() == 0
